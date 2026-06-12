"""
Nepali Devanagari OCR Pipeline
==============================
High-accuracy, high-throughput PDF → DOCX converter using GPT-4o vision.
Optimized for RAG ingestion into ChromaDB.

Features:
  - Image preprocessing (grayscale, contrast, sharpening) for maximum accuracy
  - Concurrent page-level and PDF-level parallelism
  - Automatic retry with exponential backoff
  - Resume support: skips already-processed PDFs (use --force to re-process)
  - RAG-optimized output: continuous text, markdown tables, no page separators
  - Multi-page table header preservation
  - Progress bars and detailed summary report
  - CLI arguments for all key settings

Usage:
  python ocr_pipeline.py
  python ocr_pipeline.py --dpi 220 --api-parallel 1 --api-min-interval 5 --force
  python ocr_pipeline.py --input ./my_pdfs --output ./results
"""

import os
import io
import sys
import time
import base64
import argparse
import threading
import random
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed

from dotenv import load_dotenv
from openai import OpenAI
import fitz  # PyMuPDF
from PIL import Image, ImageEnhance, ImageFilter
from docx import Document

# ─── Configuration ────────────────────────────────────────────────────────────

load_dotenv()
API_KEY = os.getenv("FREEMODEL_API_KEY")
DEFAULT_REQUEST_TIMEOUT = float(os.getenv("FREEMODEL_REQUEST_TIMEOUT", "180"))

if not API_KEY:
    print("❌ Error: FREEMODEL_API_KEY not found in .env file.")
    sys.exit(1)

client = OpenAI(
    api_key=API_KEY,
    base_url="https://api.freemodel.dev/v1",
    timeout=DEFAULT_REQUEST_TIMEOUT,
)

# ─── Prompt ───────────────────────────────────────────────────────────────────

OCR_PROMPT = (
    "Precise Nepali legal OCR. Extract only visible body text from this page in Nepali "
    "Unicode Devanagari; correct obvious OCR errors but preserve original legal wording.\n\n"
    "Output only corrected OCR text. No intro, comments, JSON, explanations, translation, "
    "summaries, paraphrase, modernization, invented text, repeated text, headers, footers, "
    "page numbers, running titles, watermarks, borders, scan noise, decorative spacing, or "
    "page separators. Keep output similar length to visible body text.\n\n"
    "Fix obvious Devanagari OCR errors: bad characters, matras, conjuncts, garbled Unicode, "
    "and common confusions ण↔न, ष↔स, ब↔व, भ↔म, घ↔प, छ↔इ. Preserve punctuation, legal terms, "
    "spelling, and visible text when uncertain.\n\n"
    "Preserve all legal numbering in Nepali Unicode exactly: १., २., ३., (१), (२), (क), "
    "(ख). Never convert to English numerals. Preserve/reconstruct visible hierarchy: "
    "दफा/धारा = dotted number, उपदफा/उपधारा = bracketed number, खण्ड = bracketed Nepali "
    "letter. Start each दफा, धारा, उपदफा, उपधारा, and खण्ड on a new line; keep simple "
    "nesting, not PDF visual spacing. Preserve body headers: भाग, अध्याय, अनुसूची, दफा, "
    "धारा, सूची, शीर्षक, नियम, उपनियम, परिच्छेद, उपदफा, उपधारा, स्पष्टीकरण.\n\n"
    "Reflow PDF-wrapped lines inside the same paragraph/clause into one continuous line. "
    "Use line breaks only at real paragraph, section, clause, heading, or list-item boundaries.\n\n"
    "Tables: always convert to Markdown using | column | separators; repeat headers for "
    "continued tables; do not preserve ASCII/visual spacing. Charts/diagrams: describe as "
    "structured Nepali text. Markdown is allowed only for tables."
)

# ─── Thread-safe progress counter ────────────────────────────────────────────

class ProgressTracker:
    """Thread-safe progress tracker with live percentage display."""

    def __init__(self, total, label="Progress"):
        self._lock = threading.Lock()
        self._done = 0
        self._failed = 0
        self._total = total
        self._label = label

    def complete(self, success=True):
        with self._lock:
            self._done += 1
            if not success:
                self._failed += 1
            pct = (self._done / self._total) * 100
            bar_len = 30
            filled = int(bar_len * self._done / self._total)
            bar = "█" * filled + "░" * (bar_len - filled)
            sys.stdout.write(
                f"\r   {self._label}: {bar} {self._done}/{self._total} ({pct:.0f}%)"
            )
            sys.stdout.flush()
            if self._done == self._total:
                sys.stdout.write("\n")
                sys.stdout.flush()

    @property
    def stats(self):
        with self._lock:
            return self._done, self._failed


class ApiThrottle:
    """Limit global API concurrency and pace request starts across all workers."""

    def __init__(self, max_concurrent=2, min_interval=2.0):
        self._semaphore = threading.BoundedSemaphore(max(1, max_concurrent))
        self._min_interval = max(0.0, min_interval)
        self._lock = threading.Lock()
        self._next_request_at = 0.0

    def acquire(self):
        self._semaphore.acquire()
        with self._lock:
            now = time.monotonic()
            if now < self._next_request_at:
                time.sleep(self._next_request_at - now)
                now = time.monotonic()
            self._next_request_at = now + self._min_interval

    def release(self):
        self._semaphore.release()


# ─── Image Preprocessing ─────────────────────────────────────────────────────

def preprocess_image(png_bytes, image_format="jpeg", jpeg_quality=85):
    """
    Enhance a raw PDF page image for maximum OCR accuracy.
    Steps:
      1. Convert to grayscale (removes color noise, reduces payload)
      2. Boost contrast (makes text stand out from background)
      3. Sharpen (crisper character edges for Devanagari strokes)
      4. Export as optimized JPEG or PNG
    """
    img = Image.open(io.BytesIO(png_bytes))
    img = img.convert("L")
    img = ImageEnhance.Contrast(img).enhance(1.5)
    img = img.filter(ImageFilter.SHARPEN)

    out = io.BytesIO()
    if image_format.lower() == "png":
        img.save(out, format="PNG", optimize=True)
        return out.getvalue(), "image/png"

    img.save(
        out,
        format="JPEG",
        quality=max(1, min(95, jpeg_quality)),
        optimize=True,
    )
    return out.getvalue(), "image/jpeg"


# ─── API Call ─────────────────────────────────────────────────────────────────

def perform_ocr_on_page(
    base64_image,
    image_mime_type,
    page_label="",
    max_retries=3,
    throttle=None,
    retry_base_delay=5.0,
    retry_max_delay=60.0,
    request_timeout=DEFAULT_REQUEST_TIMEOUT,
    image_detail="auto",
):
    """Sends the page image to GPT-4o with retry + exponential backoff."""
    payload_mb = len(base64_image) * 3 / 4 / (1024 * 1024)
    for attempt in range(1, max_retries + 1):
        try:
            if throttle:
                throttle.acquire()
            try:
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {"type": "text", "text": OCR_PROMPT},
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:{image_mime_type};base64,{base64_image}",
                                        "detail": image_detail
                                    }
                                }
                            ]
                        }
                    ],
                    max_tokens=4096,
                    temperature=0.0,
                    timeout=request_timeout,
                )
            finally:
                if throttle:
                    throttle.release()
            return response.choices[0].message.content
        except Exception as e:
            if attempt < max_retries:
                error_text = str(e).lower()
                rate_limited = "rate limit" in error_text or "429" in error_text
                timed_out = "timed out" in error_text or "timeout" in error_text
                wait = min(retry_max_delay, retry_base_delay * (2 ** (attempt - 1)))
                if rate_limited or timed_out:
                    wait = min(retry_max_delay, wait * 2)
                wait += random.uniform(0, 1.5)
                print(
                    f"\n   ⚠ {page_label} attempt {attempt}/{max_retries} failed: {e}. "
                    f"Payload ~{payload_mb:.1f} MB. Retrying in {wait:.1f}s..."
                )
                time.sleep(wait)
            else:
                print(f"\n   ✗ {page_label} FAILED after {max_retries} attempts: {e}")
                return ""


# ─── Single Page Worker ──────────────────────────────────────────────────────

def ocr_single_page(pdf_path, page_index, total_pages, dpi, max_retries, tracker, args):
    """Render one PDF page, preprocess it, OCR it. Fully thread-safe."""
    page_num = page_index + 1
    label = f"[p{page_num}/{total_pages}]"

    pdf_doc = fitz.open(pdf_path)
    page = pdf_doc.load_page(page_index)
    pix = page.get_pixmap(dpi=dpi)
    raw_png = pix.tobytes("png")
    pdf_doc.close()

    enhanced_image, image_mime_type = preprocess_image(
        raw_png,
        image_format=getattr(args, "image_format", "jpeg"),
        jpeg_quality=getattr(args, "jpeg_quality", 85),
    )
    base64_img = base64.b64encode(enhanced_image).decode("utf-8")

    text = perform_ocr_on_page(
        base64_img,
        image_mime_type,
        page_label=label,
        max_retries=max_retries,
        throttle=args.api_throttle,
        retry_base_delay=getattr(args, "retry_base_delay", 5.0),
        retry_max_delay=getattr(args, "retry_max_delay", 60.0),
        request_timeout=getattr(args, "request_timeout", DEFAULT_REQUEST_TIMEOUT),
        image_detail=getattr(args, "image_detail", "auto"),
    )

    tracker.complete(success=bool(text))
    return page_num, text


def add_text_as_docx_paragraphs(doc, text):
    """Write hard paragraph breaks to DOCX; skip empty lines to ensure single line breaks."""
    for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        if line.strip():
            doc.add_paragraph(line.rstrip())


# ─── Single PDF Processor ────────────────────────────────────────────────────

def process_single_pdf(pdf_path, args):
    """
    Process one PDF end-to-end:
      1. Check if already processed (skip unless --force)
      2. Render + preprocess pages in parallel
      3. Save .docx output (continuous text, no page separators)
    """
    output_folder = Path(args.output)
    docx_path = output_folder / f"{pdf_path.stem}.docx"
    if not hasattr(args, "api_throttle"):
        args.api_throttle = ApiThrottle(
            getattr(args, "api_parallel", 2),
            getattr(args, "api_min_interval", 2.0),
        )

    # Skip if already processed
    if not args.force and docx_path.exists():
        print(f"\n⏭  Skipping (already exists): {pdf_path.name}")
        return {"file": pdf_path.name, "status": "skipped", "pages": 0, "failed": 0, "time": 0}

    print(f"\n{'─'*60}")
    print(f"📄 {pdf_path.name}")
    print(f"{'─'*60}")

    start = time.time()

    try:
        pdf_doc = fitz.open(pdf_path)
        total_pages = len(pdf_doc)
        pdf_doc.close()

        tracker = ProgressTracker(total_pages, label="OCR")

        # OCR all pages concurrently
        results = {}
        with ThreadPoolExecutor(max_workers=args.pages_parallel) as executor:
            futures = {
                executor.submit(
                    ocr_single_page,
                    str(pdf_path), i, total_pages, args.dpi, args.retries, tracker, args
                ): i
                for i in range(total_pages)
            }
            for future in as_completed(futures):
                page_num, text = future.result()
                results[page_num] = text

        # Build .docx — continuous body text only, no generated headings or separators
        doc = Document()

        merged_text = "\n".join(
            results[pn] for pn in sorted(results.keys()) if results[pn]
        )

        if merged_text:
            add_text_as_docx_paragraphs(doc, merged_text)
        else:
            doc.add_paragraph("[⚠ Failed to extract any text from this document]")

        doc.save(docx_path)

        elapsed = time.time() - start
        _, failed = tracker.stats
        print(f"   ✅ Saved: {docx_path.name}  ({total_pages} pages, {elapsed:.1f}s)")

        return {
            "file": pdf_path.name,
            "status": "success",
            "pages": total_pages,
            "failed": failed,
            "time": elapsed
        }

    except Exception as e:
        elapsed = time.time() - start
        print(f"   ❌ FAILED: {e}")
        return {"file": pdf_path.name, "status": "error", "pages": 0, "failed": 0, "time": elapsed, "error": str(e)}


# ─── Main Orchestrator ───────────────────────────────────────────────────────

def process_pdfs(args):
    """Find all PDFs and process them with maximum throughput."""
    input_folder = Path(args.input)
    output_folder = Path(args.output)
    input_folder.mkdir(exist_ok=True)
    output_folder.mkdir(exist_ok=True)
    args.api_throttle = ApiThrottle(args.api_parallel, args.api_min_interval)

    pdf_files = sorted(input_folder.glob("*.pdf"))

    if not pdf_files:
        print(f"No PDF files found in '{input_folder}/'. Drop some files there and retry.")
        return

    total_start = time.time()

    print("╔════════════════════════════════════════════════════════════╗")
    print("║          Nepali Devanagari OCR Pipeline                   ║")
    print("╚════════════════════════════════════════════════════════════╝")
    print(f"  PDFs found      : {len(pdf_files)}")
    print(f"  DPI              : {args.dpi}")
    print(f"  Page concurrency : {args.pages_parallel} per PDF")
    print(f"  PDF concurrency  : {args.pdfs_parallel}")
    print(f"  API concurrency  : {args.api_parallel} total")
    print(f"  API start delay  : {args.api_min_interval:.1f}s minimum")
    print(f"  Request timeout  : {args.request_timeout:.0f}s")
    print(f"  Image format     : {args.image_format.upper()}")
    if args.image_format == "jpeg":
        print(f"  JPEG quality     : {args.jpeg_quality}")
    print(f"  Image detail     : {args.image_detail}")
    print(f"  Retries          : {args.retries}")
    print(f"  Force re-process : {'Yes' if args.force else 'No (skips existing)'}")
    print(f"  Output folder    : {output_folder.resolve()}")

    # Process PDFs
    reports = []
    if len(pdf_files) == 1 or args.pdfs_parallel == 1:
        for pdf_path in pdf_files:
            reports.append(process_single_pdf(pdf_path, args))
    else:
        with ThreadPoolExecutor(max_workers=args.pdfs_parallel) as executor:
            futures = {executor.submit(process_single_pdf, p, args): p for p in pdf_files}
            for future in as_completed(futures):
                reports.append(future.result())

    # ── Summary Report ────────────────────────────────────────────────────
    total_elapsed = time.time() - total_start

    succeeded = [r for r in reports if r["status"] == "success"]
    skipped = [r for r in reports if r["status"] == "skipped"]
    failed = [r for r in reports if r["status"] == "error"]
    total_pages = sum(r["pages"] for r in reports)
    total_failed_pages = sum(r["failed"] for r in reports)

    print(f"\n╔════════════════════════════════════════════════════════════╗")
    print(f"║                    Summary Report                        ║")
    print(f"╠════════════════════════════════════════════════════════════╣")
    print(f"║  Total PDFs     : {len(pdf_files):<39}║")
    print(f"║  ✅ Succeeded   : {len(succeeded):<39}║")
    print(f"║  ⏭  Skipped     : {len(skipped):<39}║")
    print(f"║  ❌ Failed      : {len(failed):<39}║")
    print(f"║  Total pages    : {total_pages:<39}║")
    print(f"║  Failed pages   : {total_failed_pages:<39}║")
    print(f"║  Total time     : {total_elapsed:.1f}s{'':<35}║")
    if total_pages > 0:
        avg = total_elapsed / total_pages
        print(f"║  Avg per page   : {avg:.1f}s{'':<35}║")
    print(f"╚════════════════════════════════════════════════════════════╝")

    if failed:
        print("\n⚠ Failed files:")
        for r in failed:
            print(f"   - {r['file']}: {r.get('error', 'Unknown error')}")

    if skipped:
        print(f"\n💡 {len(skipped)} file(s) skipped (already processed). Use --force to re-process.")


# ─── CLI ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Nepali Devanagari OCR Pipeline — PDF to DOCX using GPT-4o vision (RAG-optimized)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python ocr_pipeline.py                           # Process all PDFs in ./data
  python ocr_pipeline.py --force                    # Re-process even if output exists
  python ocr_pipeline.py --dpi 220 --api-parallel 1 --api-min-interval 5
  python ocr_pipeline.py --input ./my_pdfs --output ./results
        """
    )
    parser.add_argument("--input", default="data", help="Input folder with PDFs (default: data)")
    parser.add_argument("--output", default="output", help="Output folder for results (default: output)")
    parser.add_argument("--dpi", type=int, default=220, help="Image render DPI (default: 220)")
    parser.add_argument("--pages-parallel", type=int, default=2, help="Concurrent pages per PDF (default: 2)")
    parser.add_argument("--pdfs-parallel", type=int, default=1, help="Concurrent PDFs (default: 1)")
    parser.add_argument("--api-parallel", type=int, default=1, help="Maximum total concurrent API calls (default: 1)")
    parser.add_argument("--api-min-interval", type=float, default=5.0, help="Minimum seconds between API request starts (default: 5.0)")
    parser.add_argument("--request-timeout", type=float, default=DEFAULT_REQUEST_TIMEOUT, help="API request timeout in seconds (default: 180 or FREEMODEL_REQUEST_TIMEOUT)")
    parser.add_argument("--image-format", choices=["jpeg", "png"], default="jpeg", help="Encoded page image format sent to the API (default: jpeg)")
    parser.add_argument("--jpeg-quality", type=int, default=85, help="JPEG quality when --image-format jpeg is used (default: 85)")
    parser.add_argument("--image-detail", choices=["auto", "low", "high"], default="auto", help="Vision detail hint sent to the API (default: auto)")
    parser.add_argument("--retries", type=int, default=3, help="API retry attempts per page (default: 3)")
    parser.add_argument("--retry-base-delay", type=float, default=5.0, help="Initial retry delay in seconds (default: 5.0)")
    parser.add_argument("--retry-max-delay", type=float, default=60.0, help="Maximum retry delay in seconds (default: 60.0)")
    parser.add_argument("--force", action="store_true", help="Re-process PDFs even if output already exists")

    args = parser.parse_args()
    process_pdfs(args)


if __name__ == "__main__":
    main()
